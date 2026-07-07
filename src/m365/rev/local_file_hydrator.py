"""REV LocalFileHydrator (Zone C) — FR-PCI-5, Phase 3 (P3-5).

specs/gaps.md P3-5. Hydrates locally-downloaded ``.docx`` and ``.pdf`` files
into ``HydratedContent`` objects for the REV extraction pipeline.

**Word (.docx) extraction:**
* Requires ``python-docx>=1.1`` (``pip install python-docx``). If absent,
  hydration returns ``Unsupported(python_docx_not_installed)``.
* Macro/VBA denial: checks for ``word/vbaProject.bin`` in the ZIP before
  opening with python-docx. Files containing embedded VBA are quarantined with
  ``reason=macro_denied`` (tracked via ``macro_denied_count`` counter).
* Extracts text from all paragraphs and tables. Empty documents → ``metadata_only``.
* Embedded objects (OLE, images) are silently skipped by python-docx.

**PDF (.pdf) extraction:**
* Uses ``pypdf>=5.7.0`` (already in core dependencies).
* ``pdf_no_text``: if all pages yield empty text, returns ``Unsupported`` with
  ``reason=pdf_no_text`` so the pipeline can quarantine the file (scanned
  image PDFs with no embedded text layer).
* Encrypted PDFs: returns ``Unsupported(pdf_encrypted)``.

**Common:**
* ``EnumeratedCandidate.partial_metadata["file_path"]`` carries the claimed
  file path; ``partial_metadata["file_sha256"]`` is the logical dedup key.
* 10 MB size guard (defensive backstop; LocalFileEnumerator quarantines at
  claim time).
* 30s per-file timeout on POSIX; no hard timeout on Windows (local file).
* All extraction errors → ``Unsupported(hydration_error: ...)`` rather than
  exceptions, so the pipeline can quarantine the file gracefully.

Per-cycle counters (incremented during ``hydrate`` calls):
  * ``macro_denied_count`` — Word files quarantined for containing VBA macros
  * ``pdf_no_text_count`` — PDF files with no extractable text

Zone C: imports only ``src.core.*`` + stdlib + python-docx + pypdf.
"""

from __future__ import annotations

import logging
import signal
import sys
import zipfile
from collections.abc import Callable
from datetime import datetime, timezone
from pathlib import Path

from src.core.rev.entity_types import EntityType
from src.core.rev.identity import CanonicalItemIdentity
from src.core.rev.normalizer import (
    chunk_canonical,
    normalize_whitespace,
    normalized_source_hash,
)
from src.core.rev.ports import EnumeratedCandidate, HydratedContent
from src.core.rev.result import PortResult, Success, Unsupported

log = logging.getLogger(__name__)

_TIMEOUT_SECONDS = 30
_MIN_BODY_CHARS = 10
_MAX_FILE_BYTES = 10 * 1024 * 1024   # 10 MB per file
_VBA_ENTRY = "word/vbaProject.bin"   # macro indicator in .docx ZIP

try:
    import docx as _docx_module  # type: ignore[import-untyped]
    _DOCX_AVAILABLE = True
except ImportError:
    _docx_module = None  # type: ignore[assignment]
    _DOCX_AVAILABLE = False

# pypdf is a core dependency; no import guard needed.
try:
    from pypdf import PdfReader as _PdfReader  # type: ignore[import-untyped]
    _PYPDF_AVAILABLE = True
except ImportError:
    _PdfReader = None  # type: ignore[assignment,misc]
    _PYPDF_AVAILABLE = False


class _TimeoutError(Exception):
    pass


def _alarm_handler(signum: int, frame: object) -> None:
    raise _TimeoutError("LocalFileHydrator: per-file timeout exceeded")


def _has_vba_macros(path: Path) -> bool:
    """Return True if a .docx ZIP contains a VBA project (macro indicator)."""
    try:
        with zipfile.ZipFile(path, "r") as zf:
            return _VBA_ENTRY in zf.namelist()
    except Exception:
        return False


def _extract_docx_text(path: Path) -> str:
    """Extract all paragraph + table text from a .docx file using python-docx."""
    doc = _docx_module.Document(str(path))
    parts: list[str] = []

    for para in doc.paragraphs:
        text = para.text.strip()
        if text:
            parts.append(text)

    for table in doc.tables:
        for row in table.rows:
            row_texts = [cell.text.strip() for cell in row.cells if cell.text.strip()]
            if row_texts:
                parts.append(" | ".join(row_texts))

    return "\n".join(parts)


def _extract_pdf_text(path: Path) -> str | None:
    """Extract text from all pages of a PDF.

    Returns None if no text could be extracted (scanned image PDF) or if the
    PDF is encrypted. Raises for parse errors.
    """
    reader = _PdfReader(str(path))
    if reader.is_encrypted:
        raise ValueError("pdf_encrypted")
    pages_text: list[str] = []
    for page in reader.pages:
        text = page.extract_text() or ""
        if text.strip():
            pages_text.append(text)
    if not pages_text:
        return None
    return "\n".join(pages_text)


class LocalFileHydrator:
    """Hydrate a locally-claimed ``.docx`` or ``.pdf`` file → ``HydratedContent``.

    Satisfies the ``ContentHydrator`` Protocol (Zone A ``src/core/rev/ports.py``).
    ``EnumeratedCandidate.partial_metadata["file_path"]`` carries the claimed
    file path.

    Per-cycle counters:
    * ``macro_denied_count`` — Word files quarantined for containing VBA macros
    * ``pdf_no_text_count`` — PDF files returning ``Unsupported(pdf_no_text)``

    Usage::

        hydrator = LocalFileHydrator(
            mailbox_tenant_id="tenant-abc",
            principal_mailbox="tpm@example.com",
        )
        result = hydrator.hydrate(candidate, correlation_id="rev-cycle-docs-001")
    """

    def __init__(
        self,
        *,
        mailbox_tenant_id: str,
        principal_mailbox: str,
        container: str = "docs",
        timeout_seconds: int = _TIMEOUT_SECONDS,
    ) -> None:
        self._tenant_id = mailbox_tenant_id
        self._principal_mailbox = principal_mailbox
        self._container = container
        self._timeout_seconds = timeout_seconds
        # Per-cycle counters
        self.macro_denied_count: int = 0
        self.pdf_no_text_count: int = 0

    def hydrate(
        self,
        candidate: EnumeratedCandidate,
        *,
        correlation_id: str,
    ) -> PortResult[HydratedContent]:
        """Hydrate one .docx or .pdf file. Returns Unsupported if file_path is absent."""
        file_path_str = candidate.partial_metadata.get("file_path")
        if not file_path_str:
            return Unsupported(
                entity_type=EntityType.DRIVE_ITEM.value,
                reason="file_path_missing: candidate has no file_path in partial_metadata",
            )
        file_path = Path(str(file_path_str))
        if not file_path.exists():
            return Unsupported(
                entity_type=EntityType.DRIVE_ITEM.value,
                reason=f"file_not_found: {file_path}",
            )

        # Size guard (defensive backstop; LocalFileEnumerator quarantines at claim time).
        try:
            size = file_path.stat().st_size
        except OSError:
            size = 0
        if size > _MAX_FILE_BYTES:
            return Unsupported(
                entity_type=EntityType.DRIVE_ITEM.value,
                reason=f"size_exceeded: {size} bytes > {_MAX_FILE_BYTES}",
            )

        suffix = file_path.suffix.lower()
        if suffix == ".docx":
            return self._hydrate_docx(file_path, candidate=candidate, correlation_id=correlation_id)
        if suffix == ".pdf":
            return self._hydrate_pdf(file_path, candidate=candidate, correlation_id=correlation_id)
        return Unsupported(
            entity_type=EntityType.DRIVE_ITEM.value,
            reason=f"unsupported_format: {suffix} (only .docx and .pdf are supported)",
        )

    def _run_with_timeout(
        self,
        fn: Callable[[Path], str | None],
        path: Path,
    ) -> str | None:
        """Run a text-extraction function with per-file POSIX timeout."""
        use_timeout = sys.platform != "win32" and hasattr(signal, "SIGALRM")
        if use_timeout:
            signal.signal(signal.SIGALRM, _alarm_handler)  # type: ignore[attr-defined]
            signal.alarm(self._timeout_seconds)  # type: ignore[attr-defined]
        try:
            return fn(path)
        finally:
            if use_timeout:
                signal.alarm(0)  # type: ignore[attr-defined]
                signal.signal(signal.SIGALRM, signal.SIG_DFL)  # type: ignore[attr-defined]

    def _hydrate_docx(
        self,
        path: Path,
        *,
        candidate: EnumeratedCandidate,
        correlation_id: str,
    ) -> PortResult[HydratedContent]:
        if not _DOCX_AVAILABLE:
            return Unsupported(
                entity_type=EntityType.DRIVE_ITEM.value,
                reason="python_docx_not_installed: pip install 'python-docx>=1.1'",
            )

        # Macro denial: reject files with embedded VBA before opening.
        if _has_vba_macros(path):
            log.warning("LocalFileHydrator: macro_denied — %s contains VBA project", path)
            self.macro_denied_count += 1
            return Unsupported(
                entity_type=EntityType.DRIVE_ITEM.value,
                reason=f"macro_denied: {path.name} contains embedded VBA macros",
            )

        try:
            raw = self._run_with_timeout(_extract_docx_text, path)
        except _TimeoutError:
            return Unsupported(
                entity_type=EntityType.DRIVE_ITEM.value,
                reason=f"hydration_timeout: exceeded {self._timeout_seconds}s on {path.name}",
            )
        except Exception as exc:
            log.warning("LocalFileHydrator: docx extraction failed on %s: %s", path, exc)
            return Unsupported(
                entity_type=EntityType.DRIVE_ITEM.value,
                reason=f"hydration_error: {exc}",
            )

        return self._build_content(
            raw or "",
            path=path,
            candidate=candidate,
            correlation_id=correlation_id,
            source_format="docx",
        )

    def _hydrate_pdf(
        self,
        path: Path,
        *,
        candidate: EnumeratedCandidate,
        correlation_id: str,
    ) -> PortResult[HydratedContent]:
        if not _PYPDF_AVAILABLE:
            return Unsupported(
                entity_type=EntityType.DRIVE_ITEM.value,
                reason="pypdf_not_installed: pip install 'pypdf>=5.7'",
            )

        try:
            raw = self._run_with_timeout(_extract_pdf_text, path)
        except _TimeoutError:
            return Unsupported(
                entity_type=EntityType.DRIVE_ITEM.value,
                reason=f"hydration_timeout: exceeded {self._timeout_seconds}s on {path.name}",
            )
        except ValueError as exc:
            reason = str(exc)
            return Unsupported(
                entity_type=EntityType.DRIVE_ITEM.value,
                reason=reason,  # "pdf_encrypted" from _extract_pdf_text
            )
        except Exception as exc:
            log.warning("LocalFileHydrator: pdf extraction failed on %s: %s", path, exc)
            return Unsupported(
                entity_type=EntityType.DRIVE_ITEM.value,
                reason=f"hydration_error: {exc}",
            )

        if raw is None:
            log.info("LocalFileHydrator: pdf_no_text — %s", path)
            self.pdf_no_text_count += 1
            return Unsupported(
                entity_type=EntityType.DRIVE_ITEM.value,
                reason=f"pdf_no_text: {path.name} has no extractable text (scanned image?)",
            )

        return self._build_content(
            raw,
            path=path,
            candidate=candidate,
            correlation_id=correlation_id,
            source_format="pdf",
        )

    def _build_content(
        self,
        raw_text: str,
        *,
        path: Path,
        candidate: EnumeratedCandidate,
        correlation_id: str,
        source_format: str,
    ) -> PortResult[HydratedContent]:
        canonical = normalize_whitespace(raw_text)
        file_sha256 = str(candidate.partial_metadata.get("file_sha256", candidate.locator.resource_id))

        route_metadata: dict[str, object] = {
            "file_name": path.name,
            "file_sha256": file_sha256,
            "source_format": source_format,
        }

        identity = CanonicalItemIdentity(
            source_type=EntityType.DRIVE_ITEM,
            tenant_id=self._tenant_id,
            principal_mailbox=self._principal_mailbox,
            container=self._container,
            resource_id=file_sha256,
        )

        source_hash = normalized_source_hash(canonical if canonical else "")
        chunks = chunk_canonical(canonical) if canonical else ()
        is_metadata_only = len(canonical.replace(" ", "")) < _MIN_BODY_CHARS

        return Success(HydratedContent(
            identity=identity,
            canonical_text=canonical,
            normalized_source_hash=source_hash,
            chunks=chunks,
            route_metadata=route_metadata,
            hydration_rung="metadata_only" if is_metadata_only else f"{source_format}_body",
            metadata_only=is_metadata_only,
            retrieved_at=datetime.now(timezone.utc),
            correlation_id=correlation_id,
        ))


__all__ = ["LocalFileHydrator"]
